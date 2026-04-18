import os
import io
import uuid
import tempfile
import logging
from pathlib import Path
from flask import Flask, render_template, request, send_file, jsonify
from werkzeug.utils import secure_filename

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50 MB
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", "camargo-apps-secret")

ALLOWED_EXTENSIONS = {"pdf"}


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def _pick_runtime_tmpdir() -> str:
    """
    Escolhe um diretório temporário gravável para execução.
    Evita falhas de permissão no temp padrão do SO.
    """
    candidates = []
    env_tmp = os.environ.get("CONVERSOR_TMP_DIR", "").strip()
    if env_tmp:
        candidates.append(env_tmp)
    candidates.append(os.path.join(os.getcwd(), "_runtime_tmp"))
    candidates.append(tempfile.gettempdir())

    for base in candidates:
        try:
            os.makedirs(base, exist_ok=True)
            probe = os.path.join(base, f".probe-{uuid.uuid4().hex}")
            with open(probe, "wb") as f:
                f.write(b"ok")
            os.remove(probe)
            return base
        except Exception:
            continue
    raise RuntimeError("Nenhum diretório temporário gravável disponível.")


def _libreoffice_bin() -> str:
    """Localiza o executável do LibreOffice (Windows ou Linux)."""
    import shutil
    # No Windows, prioriza caminho absoluto conhecido para evitar binários quebrados no PATH.
    for p in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        r"C:\Program Files\LibreOffice\program\soffice.com",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.com",
    ):
        if os.path.isfile(p):
            return p
    # Fallback por PATH.
    for cmd in ("soffice.com", "libreoffice", "soffice"):
        found = shutil.which(cmd)
        if found:
            return found
    raise FileNotFoundError("LibreOffice não encontrado.")


def _prepare_lo_env(lo_bin: str) -> dict:
    """
    Prepara ambiente isolado para subprocesso do LibreOffice.
    Evita erro 'Could not find platform independent libraries' em alguns ambientes.
    """
    env = os.environ.copy()
    # Remove variáveis de Python que podem interferir no runtime interno do LibreOffice.
    for key in ("PYTHONHOME", "PYTHONPATH", "PYTHONEXECUTABLE", "VIRTUAL_ENV", "CONDA_PREFIX"):
        env.pop(key, None)

    lo_dir = os.path.dirname(os.path.abspath(lo_bin))
    if lo_dir:
        env["PATH"] = lo_dir + os.pathsep + env.get("PATH", "")
        env["URE_BOOTSTRAP"] = f"vnd.sun.star.pathname:{os.path.join(lo_dir, 'fundamental.ini')}"

        # Alguns ambientes Windows precisam de PYTHONHOME explícito do runtime embutido do LO.
        try:
            pycore = next(
                (
                    os.path.join(lo_dir, d)
                    for d in os.listdir(lo_dir)
                    if d.lower().startswith("python-core-") and os.path.isdir(os.path.join(lo_dir, d))
                ),
                None,
            )
            if pycore:
                env["PYTHONHOME"] = pycore
        except Exception:
            pass
    env["PYTHONNOUSERSITE"] = "1"

    env["SAL_DISABLE_SYNCHRONOUS_PRINTER_DETECTION"] = "1"
    env["SAL_DISABLE_PRINTERLIST"] = "1"
    env["SAL_USE_VCLPLUGIN"] = "svp"
    return env


def _docx_via_libreoffice(pdf_path: str, output_path: str) -> None:
    """Converte via LibreOffice headless — MPL-2.0, uso comercial livre."""
    import subprocess, shutil, tempfile
    lo = _libreoffice_bin()
    out_dir = os.path.dirname(output_path)
    stem = Path(pdf_path).stem
    lo_out = os.path.join(out_dir, f"{stem}.docx")

    def _run_convert(input_path: str, convert_to: str, outdir: str, timeout_s: int,
                     infilter: str | None = None, profile_uri: str | None = None) -> tuple[bool, str]:
        cmd = [
            lo,
            "--headless",
            "--invisible",
            "--norestore",
            "--nodefault",
            "--nolockcheck",
            "--nofirststartwizard",
            "--nologo",
        ]
        if profile_uri:
            cmd.append(f"-env:UserInstallation={profile_uri}")
        if infilter:
            cmd.append(f"--infilter={infilter}")
        cmd.extend(["--convert-to", convert_to, "--outdir", outdir, input_path])
        try:
            run_env = _prepare_lo_env(lo)
            lo_cwd = os.path.dirname(os.path.abspath(lo))
            startupinfo = None
            creationflags = 0
            if os.name == "nt":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                startupinfo.wShowWindow = 0
                creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
            r = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=timeout_s,
                env=run_env,
                cwd=lo_cwd,
                startupinfo=startupinfo,
                creationflags=creationflags,
            )
            if r.returncode == 0:
                return True, ""
            msg = (r.stderr or "").strip() or (r.stdout or "").strip() or "erro sem detalhes"
            return False, msg
        except subprocess.TimeoutExpired as te:
            return False, f"timeout ({timeout_s}s)"
        except Exception as ex:
            return False, str(ex)

    def _ensure_docx(at_path: str) -> bool:
        return os.path.isfile(at_path) and os.path.getsize(at_path) > 0

    # Sempre usa perfil isolado: evita anexar a sessão aberta do LibreOffice (popup de impressora).
    lo_profile = tempfile.mkdtemp(prefix="lo-profile-", dir=_pick_runtime_tmpdir())
    profile_uri = Path(lo_profile).as_uri()
    errors: list[str] = []
    try:
        # 1) Pipeline mais fiel: PDF -> ODT -> DOCX.
        tdir = tempfile.mkdtemp(prefix="lo-odt-", dir=_pick_runtime_tmpdir())
        try:
            odt_path = os.path.join(tdir, f"{stem}.odt")
            ok, err = _run_convert(
                input_path=pdf_path,
                convert_to="odt:writer8",
                outdir=tdir,
                timeout_s=90,
                infilter="draw_pdf_import",
                profile_uri=profile_uri,
            )
            if ok and os.path.isfile(odt_path):
                ok2, err2 = _run_convert(
                    input_path=odt_path,
                    convert_to="docx:MS Word 2007 XML",
                    outdir=out_dir,
                    timeout_s=90,
                    profile_uri=profile_uri,
                )
                if ok2 and _ensure_docx(lo_out):
                    if os.path.abspath(lo_out) != os.path.abspath(output_path):
                        shutil.move(lo_out, output_path)
                    return
                errors.append(f"ODT->DOCX: {err2}")
            else:
                errors.append(f"PDF->ODT: {err}")
        finally:
            shutil.rmtree(tdir, ignore_errors=True)

        # 2) Draw direto -> DOCX (fallback).
        ok, err = _run_convert(
            input_path=pdf_path,
            convert_to="docx:MS Word 2007 XML",
            outdir=out_dir,
            timeout_s=90,
            infilter="draw_pdf_import",
            profile_uri=profile_uri,
        )
        if ok and _ensure_docx(lo_out):
            if os.path.abspath(lo_out) != os.path.abspath(output_path):
                shutil.move(lo_out, output_path)
            return
        errors.append(f"PDF->DOCX (Draw): {err}")

        # 3) Fallback: Writer import padrão -> DOCX.
        ok, err = _run_convert(
            input_path=pdf_path,
            convert_to="docx:MS Word 2007 XML",
            outdir=out_dir,
            timeout_s=60,
            profile_uri=profile_uri,
        )
        if ok and _ensure_docx(lo_out):
            if os.path.abspath(lo_out) != os.path.abspath(output_path):
                shutil.move(lo_out, output_path)
            return
        errors.append(f"PDF->DOCX (padrão): {err}")

        # 4) Tentativa em 2 etapas (Draw): PDF -> ODG -> DOCX.
        tdir = tempfile.mkdtemp(prefix="lo-draw-", dir=_pick_runtime_tmpdir())
        try:
            odg_path = os.path.join(tdir, f"{stem}.odg")
            ok, err = _run_convert(
                input_path=pdf_path,
                convert_to="odg",
                outdir=tdir,
                timeout_s=60,
                infilter="draw_pdf_import",
                profile_uri=profile_uri,
            )
            if ok and os.path.isfile(odg_path):
                ok2, err2 = _run_convert(
                    input_path=odg_path,
                    convert_to="docx:MS Word 2007 XML",
                    outdir=out_dir,
                    timeout_s=60,
                    profile_uri=profile_uri,
                )
                if ok2 and _ensure_docx(lo_out):
                    if os.path.abspath(lo_out) != os.path.abspath(output_path):
                        shutil.move(lo_out, output_path)
                    return
                errors.append(f"ODG->DOCX: {err2}")
            else:
                errors.append(f"PDF->ODG: {err}")
        finally:
            shutil.rmtree(tdir, ignore_errors=True)
    finally:
        shutil.rmtree(lo_profile, ignore_errors=True)

    raise RuntimeError(" | ".join(e for e in errors if e) or "Erro no LibreOffice")


def _odt_via_libreoffice(pdf_path: str, output_path: str) -> None:
    """Converte PDF para ODT via LibreOffice (modo Draw, perfil isolado)."""
    import subprocess, shutil, tempfile
    lo = _libreoffice_bin()
    out_dir = os.path.dirname(output_path)
    stem = Path(pdf_path).stem
    lo_out = os.path.join(out_dir, f"{stem}.odt")

    def _run_convert(input_path: str, convert_to: str, timeout_s: int,
                     infilter: str | None = None, profile_uri: str | None = None) -> tuple[bool, str]:
        cmd = [
            lo,
            "--headless",
            "--invisible",
            "--norestore",
            "--nodefault",
            "--nolockcheck",
            "--nofirststartwizard",
            "--nologo",
        ]
        if profile_uri:
            cmd.append(f"-env:UserInstallation={profile_uri}")
        if infilter:
            cmd.append(f"--infilter={infilter}")
        cmd.extend(["--convert-to", convert_to, "--outdir", out_dir, input_path])
        try:
            run_env = _prepare_lo_env(lo)
            lo_cwd = os.path.dirname(os.path.abspath(lo))
            startupinfo = None
            creationflags = 0
            if os.name == "nt":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                startupinfo.wShowWindow = 0
                creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
            r = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=timeout_s,
                env=run_env,
                cwd=lo_cwd,
                startupinfo=startupinfo,
                creationflags=creationflags,
            )
            if r.returncode == 0:
                return True, ""
            msg = (r.stderr or "").strip() or (r.stdout or "").strip() or "erro sem detalhes"
            return False, msg
        except subprocess.TimeoutExpired:
            return False, f"timeout ({timeout_s}s)"
        except Exception as ex:
            return False, str(ex)

    lo_profile = tempfile.mkdtemp(prefix="lo-profile-", dir=_pick_runtime_tmpdir())
    profile_uri = Path(lo_profile).as_uri()
    errors: list[str] = []
    try:
        ok, err = _run_convert(
            input_path=pdf_path,
            convert_to="odt:writer8",
            timeout_s=180,
            profile_uri=profile_uri,
        )
        if ok and os.path.isfile(lo_out) and os.path.getsize(lo_out) > 0:
            if os.path.abspath(lo_out) != os.path.abspath(output_path):
                shutil.move(lo_out, output_path)
            return
        errors.append(f"PDF->ODT: {err}")

        # Fallback com importador Draw explícito.
        ok, err = _run_convert(
            input_path=pdf_path,
            convert_to="odt:writer8",
            timeout_s=180,
            infilter="draw_pdf_import",
            profile_uri=profile_uri,
        )
        if ok and os.path.isfile(lo_out) and os.path.getsize(lo_out) > 0:
            if os.path.abspath(lo_out) != os.path.abspath(output_path):
                shutil.move(lo_out, output_path)
            return
        errors.append(f"PDF->ODT (Draw): {err}")

    finally:
        shutil.rmtree(lo_profile, ignore_errors=True)

    raise RuntimeError(" | ".join(e for e in errors if e) or "Erro no LibreOffice")


def _odt_to_docx_via_libreoffice(odt_path: str, output_path: str) -> None:
    """Converte ODT -> DOCX via LibreOffice com perfil isolado."""
    import subprocess, shutil, tempfile
    lo = _libreoffice_bin()
    out_dir = os.path.dirname(output_path)
    stem = Path(odt_path).stem
    lo_out = os.path.join(out_dir, f"{stem}.docx")

    lo_profile = tempfile.mkdtemp(prefix="lo-profile-", dir=_pick_runtime_tmpdir())
    try:
        profile_uri = Path(lo_profile).as_uri()
        run_env = _prepare_lo_env(lo)
        lo_cwd = os.path.dirname(os.path.abspath(lo))
        startupinfo = None
        creationflags = 0
        if os.name == "nt":
            startupinfo = subprocess.STARTUPINFO()
            startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            startupinfo.wShowWindow = 0
            creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)

        errors: list[str] = []
        for conv in ("docx", "docx:MS Word 2007 XML"):
            cmd = [
                lo,
                "--headless",
                "--invisible",
                "--norestore",
                "--nodefault",
                "--nolockcheck",
                "--nofirststartwizard",
                "--nologo",
                f"-env:UserInstallation={profile_uri}",
                "--convert-to", conv,
                "--outdir", out_dir,
                odt_path,
            ]
            try:
                r = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=180,
                    env=run_env,
                    cwd=lo_cwd,
                    startupinfo=startupinfo,
                    creationflags=creationflags,
                )
            except subprocess.TimeoutExpired:
                errors.append(f"ODT->DOCX ({conv}): timeout (180s)")
                continue

            if r.returncode == 0 and os.path.isfile(lo_out) and os.path.getsize(lo_out) > 0:
                if os.path.abspath(lo_out) != os.path.abspath(output_path):
                    shutil.move(lo_out, output_path)
                return

            msg = (r.stderr or "").strip() or (r.stdout or "").strip() or "erro sem detalhes"
            errors.append(f"ODT->DOCX ({conv}): {msg}")

        raise RuntimeError(" | ".join(errors) if errors else "ODT->DOCX não gerou arquivo válido.")

    finally:
        shutil.rmtree(lo_profile, ignore_errors=True)


def _odt_to_docx_via_word_com(odt_path: str, output_path: str) -> None:
    """
    Converte ODT -> DOCX usando automação do Microsoft Word (Windows).
    Usado como alternativa quando o filtro do LibreOffice falha.
    """
    import subprocess
    if os.name != "nt":
        raise RuntimeError("Conversão via Word COM disponível apenas no Windows.")

    odt_abs = str(Path(odt_path).resolve())
    docx_abs = str(Path(output_path).resolve())
    odt_ps = odt_abs.replace("'", "''")
    docx_ps = docx_abs.replace("'", "''")
    # wdFormatDocumentDefault = 16 (DOCX)
    ps_script = (
        "$ErrorActionPreference='Stop';"
        "$word=$null;$doc=$null;"
        f"$in='{odt_ps}';"
        f"$out='{docx_ps}';"
        "try {"
        "  $word = New-Object -ComObject Word.Application;"
        "  $word.Visible = $false;"
        "  $word.DisplayAlerts = 0;"
        "  $doc = $word.Documents.Open($in, $false, $true);"
        "  $doc.SaveAs([ref]$out, [ref]16);"
        "}"
        "finally {"
        "  if ($doc -ne $null) { $doc.Close([ref]$false) | Out-Null }"
        "  if ($word -ne $null) { $word.Quit() | Out-Null }"
        "}"
    )
    r = subprocess.run(
        ["powershell", "-NoProfile", "-NonInteractive", "-ExecutionPolicy", "Bypass", "-Command", ps_script],
        capture_output=True,
        text=True,
        timeout=180,
    )
    if r.returncode != 0:
        msg = (r.stderr or "").strip() or (r.stdout or "").strip() or "erro sem detalhes"
        raise RuntimeError(f"ODT->DOCX (Word COM) falhou: {msg}")
    if not os.path.isfile(docx_abs) or os.path.getsize(docx_abs) == 0:
        raise RuntimeError("ODT->DOCX (Word COM) não gerou arquivo válido.")


PT2EMU = 12700  # 1 ponto PDF = 12700 EMU (English Metric Units)


def _xml_textbox(text: str, l: int, t: int, cx: int, cy: int,
                 bold: bool, italic: bool, sz_half: int,
                 color_hex: str, sid: int) -> str:
    """XML de parágrafo com caixa de texto flutuante ancorada à página (wp:anchor)."""
    b   = "1" if bold   else "0"
    iv  = "1" if italic else "0"
    col = color_hex.lstrip("#")
    # Namespaces todos declarados na raiz para evitar conflito ao mesclar árvores lxml
    return (
        '<w:p'
        ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<w:pPr><w:spacing w:before="0" w:after="0" w:line="20" w:lineRule="exact"/></w:pPr>'
        '<w:r><w:rPr/><w:drawing>'
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"'
        f' relativeHeight="{251658240 + sid}" behindDoc="0"'
        f' locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>{l}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{t}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        f'<wp:docPr id="{sid}" name="tb{sid}"/>'
        '<wp:cNvGraphicFramePr/>'
        '<a:graphic>'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp>'
        '<wps:cNvSpPr txBox="1"><a:spLocks noChangeArrowheads="1"/></wps:cNvSpPr>'
        '<wps:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        '<a:noFill/><a:ln><a:noFill/></a:ln>'
        '</wps:spPr>'
        '<wps:txbx><w:txbxContent>'
        '<w:p>'
        '<w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr>'
        '<w:r><w:rPr>'
        f'<w:b w:val="{b}"/><w:bCs w:val="{b}"/>'
        f'<w:i w:val="{iv}"/><w:iCs w:val="{iv}"/>'
        f'<w:sz w:val="{sz_half}"/><w:szCs w:val="{sz_half}"/>'
        f'<w:color w:val="{col}"/>'
        f'</w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>'
        '</w:p>'
        '</w:txbxContent></wps:txbx>'
        '<wps:bodyPr rot="0" spcFirstLastPara="0" vertOverflow="overflow"'
        ' horzOverflow="overflow" vert="horz" wrap="none"'
        ' lIns="0" tIns="0" rIns="0" bIns="0" anchor="t" anchorCtr="0">'
        '<a:noAutofit/></wps:bodyPr>'
        '</wps:wsp></a:graphicData></a:graphic>'
        '</wp:anchor></w:drawing></w:r>'
        '</w:p>'
    )


def _xml_float_image(r_id: str, l: int, t: int, cx: int, cy: int, sid: int) -> str:
    """XML de parágrafo com imagem flutuante ancorada à página (wp:anchor)."""
    return (
        '<w:p'
        ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<w:pPr><w:spacing w:before="0" w:after="0" w:line="20" w:lineRule="exact"/></w:pPr>'
        '<w:r><w:rPr/><w:drawing>'
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"'
        f' relativeHeight="{251658241 + sid}" behindDoc="0"'
        f' locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>{l}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{t}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        f'<wp:docPr id="{sid}" name="img{sid}"/>'
        '<wp:cNvGraphicFramePr>'
        '<a:graphicFrameLocks noChangeAspect="1"/>'
        '</wp:cNvGraphicFramePr>'
        '<a:graphic>'
        '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic>'
        '<pic:nvPicPr>'
        f'<pic:cNvPr id="{sid}" name="img{sid}"/>'
        '<pic:cNvPicPr><a:picLocks noChangeAspect="1"/></pic:cNvPicPr>'
        '</pic:nvPicPr>'
        '<pic:blipFill>'
        f'<a:blip r:embed="{r_id}"/>'
        '<a:stretch><a:fillRect/></a:stretch>'
        '</pic:blipFill>'
        '<pic:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        '<a:noFill/><a:ln><a:noFill/></a:ln>'
        '</pic:spPr>'
        '</pic:pic>'
        '</a:graphicData></a:graphic>'
        '</wp:anchor></w:drawing></w:r>'
        '</w:p>'
    )


def _docx_via_reconstruction(pdf_path: str, output_path: str) -> None:
    """
    Reconstrói DOCX de forma estável com python-docx puro.
    Prioriza compatibilidade de abertura no Word (sem injeção manual de XML).
    Preserva texto, imagens e tabelas em ordem vertical por página.
    Usa: pdfplumber (MIT) + pypdf (BSD-3) + python-docx (MIT).
    """
    import pdfplumber
    from pypdf import PdfReader
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Emu
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE

    doc = Document()
    reader = PdfReader(pdf_path)
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for pg_idx, page in enumerate(pdf.pages):
                page_w_pt = float(page.width) if float(page.width) > 0 else 612.0
                page_h_pt = float(page.height) if float(page.height) > 0 else 792.0

                # Usa uma seção por página para preservar melhor dimensões/orientação.
                if pg_idx == 0:
                    sec = doc.sections[0]
                else:
                    sec = doc.add_section(WD_SECTION.NEW_PAGE)
                sec.page_width = Emu(int(page_w_pt * PT2EMU))
                sec.page_height = Emu(int(page_h_pt * PT2EMU))

                # ── Tabelas (detecção robusta) ────────────────────────
                tables = _find_tables_robust(page)
                table_bboxes = [(t["x0"], t["y"], t["x1"], t["bottom"]) for t in tables]

                # ── Linhas de texto com posição ────────────────────────
                text_lines = _page_text_lines(page, table_bboxes)

                # ── Imagens ────────────────────────────────────────────
                images = _page_images(reader.pages[pg_idx], page, page_w_pt)

                # Mescla eventos por posição vertical para manter a leitura do conteúdo
                events = (
                    [("text", t["y"], t) for t in text_lines] +
                    [("image", i["y"], i) for i in images] +
                    [("table", t["y"], t) for t in tables]
                )
                prio = {"text": 0, "table": 1, "image": 2}
                events.sort(key=lambda e: (e[1], prio.get(e[0], 9)))
                prev_y = None
                last_text_para = None
                last_text_x0 = None
                last_text_y = None
                last_text_h = None

                for kind, _, payload in events:
                    curr_y = float(payload.get("y", 0))
                    if prev_y is not None:
                        gap = curr_y - prev_y
                        if gap > 26:
                            # Mantém separação vertical aproximada entre blocos.
                            for _ in range(min(3, int(gap // 28))):
                                doc.add_paragraph("")

                    if kind == "text":
                        px0 = float(payload.get("x0", 0.0))
                        ph = float(payload.get("line_height", max(10.0, float(payload.get("size", 11.0)) * 1.2)))
                        same_para = (
                            last_text_para is not None and
                            last_text_y is not None and
                            last_text_x0 is not None and
                            abs(px0 - last_text_x0) <= 10 and
                            (curr_y - last_text_y) <= max(12.0, (last_text_h or ph) * 1.25)
                        )

                        if same_para:
                            last_text_para.add_run().add_break(WD_BREAK.LINE)
                            p = last_text_para
                            run = p.add_run(payload["text"])
                        else:
                            p = doc.add_paragraph()
                            run = p.add_run(payload["text"])

                        run.bold = payload["bold"]
                        run.italic = payload["italic"]
                        run.font.size = Pt(float(payload["size"]))
                        color = str(payload.get("color") or "").lstrip("#")
                        if len(color) == 6:
                            try:
                                run.font.color.rgb = RGBColor.from_string(color.upper())
                            except Exception:
                                pass
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.left_indent = Pt(max(0.0, px0 - 8.0))

                        x0 = float(payload.get("x0", 0.0))
                        x1 = float(payload.get("x1", page_w_pt))
                        line_w = max(0.0, x1 - x0)
                        center = (x0 + x1) / 2.0
                        if abs(center - (page_w_pt / 2.0)) < (page_w_pt * 0.08) and line_w < (page_w_pt * 0.85):
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        elif x1 > (page_w_pt * 0.92) and x0 > (page_w_pt * 0.45):
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        else:
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        last_text_para = p
                        last_text_x0 = px0
                        last_text_y = curr_y
                        last_text_h = ph
                        prev_y = curr_y
                        continue

                    if kind == "image":
                        last_text_para = None
                        data = payload["data"]
                        data.seek(0)
                        x0 = float(payload.get("x0", 0.0))
                        x1 = float(payload.get("x1", x0))
                        has_pos = bool(payload.get("has_pos", False))
                        if has_pos and x1 > x0:
                            width_pt = max(12.0, x1 - x0)
                        else:
                            # fallback aproximado quando não há bbox confiável
                            width_pt = max(18.0, min(page_w_pt * 0.9, float(payload.get("w", 120)) * 0.75))
                        width_in = max(0.14, min(7.0, width_pt / 72.0))
                        p = doc.add_paragraph()
                        p.add_run().add_picture(data, width=Inches(width_in))
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(6)
                        if has_pos:
                            p.paragraph_format.left_indent = Pt(max(0.0, x0 - 8.0))
                        prev_y = curr_y
                        continue

                    last_text_para = None
                    data = payload["data"]
                    n_rows = len(data)
                    n_cols = max((len(r) for r in data), default=1)
                    tbl = doc.add_table(rows=n_rows if n_rows > 0 else 1, cols=n_cols if n_cols > 0 else 1)
                    tbl.style = "Table Grid"
                    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                    try:
                        tbl.autofit = True
                    except Exception:
                        pass

                    col_widths_pt = payload.get("col_widths_pt") or []
                    use_fixed_widths = False
                    if len(col_widths_pt) == n_cols:
                        min_w = min(float(w) for w in col_widths_pt) if col_widths_pt else 0.0
                        max_w = max(float(w) for w in col_widths_pt) if col_widths_pt else 0.0
                        # Evita colunas absurdamente estreitas (quebra letra a letra).
                        if min_w >= 22.0 and max_w <= page_w_pt:
                            use_fixed_widths = True
                    if use_fixed_widths:
                        try:
                            tbl.autofit = False
                        except Exception:
                            pass
                        for c_idx, w_pt in enumerate(col_widths_pt):
                            try:
                                tbl.columns[c_idx].width = Pt(float(w_pt))
                            except Exception:
                                pass

                    row_heights_pt = payload.get("row_heights_pt") or []
                    for r_idx in range(n_rows):
                        row_data = data[r_idx] if r_idx < len(data) else []
                        drow = tbl.rows[r_idx]
                        if use_fixed_widths and r_idx < len(row_heights_pt):
                            try:
                                drow.height = Pt(float(row_heights_pt[r_idx]))
                                drow.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                            except Exception:
                                pass
                        for c_idx in range(n_cols):
                            val = str(row_data[c_idx] if c_idx < len(row_data) else "" or "").strip()
                            cell = drow.cells[c_idx]
                            cell.text = val
                            if r_idx == 0 and cell.paragraphs and cell.paragraphs[0].runs:
                                cell.paragraphs[0].runs[0].bold = True
                    doc.add_paragraph()
                    prev_y = float(payload.get("bottom", curr_y))

        doc.save(output_path)
    finally:
        try:
            reader.close()
        except Exception:
            pass


def convert_pdf_to_docx(pdf_path: str, output_path: str) -> None:
    """
    PDF → DOCX editável.
    1ª tentativa : LibreOffice headless (MPL-2.0) — melhor fidelidade de layout.
    Fallback      : reconstrução via pdfplumber + pypdf + python-docx (MIT/BSD).
    Sem dependências AGPL/comerciais.
    """
    try:
        _docx_via_libreoffice(pdf_path, output_path)
        logger.info("DOCX gerado via LibreOffice")
        return
    except Exception as lo_err:
        logger.warning("LibreOffice indisponível (%s) — reconstruindo via pdfplumber", lo_err)

    _docx_via_reconstruction(pdf_path, output_path)
    logger.info("DOCX gerado via reconstrução pdfplumber/pypdf")


def _configure_tesseract(pytesseract_module) -> None:
    """Configura caminho do Tesseract automaticamente quando possível."""
    import shutil

    env_cmd = os.environ.get("TESSERACT_CMD", "").strip()
    if env_cmd:
        pytesseract_module.pytesseract.tesseract_cmd = env_cmd
        return

    for cmd in ("tesseract",):
        found = shutil.which(cmd)
        if found:
            pytesseract_module.pytesseract.tesseract_cmd = found
            return

    win_candidates = (
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    )
    for path in win_candidates:
        if os.path.isfile(path):
            pytesseract_module.pytesseract.tesseract_cmd = path
            return


def _color_to_hex(color) -> str:
    """Converte cor do pdfplumber (grayscale, RGB ou CMYK) para hex."""
    if color is None or color == 0 or color == 0.0:
        return "#000000"
    if isinstance(color, (int, float)):
        v = int((1.0 - float(color)) * 255)
        return f"#{v:02X}{v:02X}{v:02X}"
    if isinstance(color, (list, tuple)):
        if len(color) == 3:
            r, g, b = (min(int(c * 255), 255) for c in color)
            return f"#{r:02X}{g:02X}{b:02X}"
        if len(color) == 4:
            c, m, y, k = color
            r = int((1 - c) * (1 - k) * 255)
            g = int((1 - m) * (1 - k) * 255)
            b = int((1 - y) * (1 - k) * 255)
            return f"#{r:02X}{g:02X}{b:02X}"
    return "#000000"


def _norm_size(size: float) -> int:
    return max(8, min(int(round(size / 2) * 2), 24))


def _compact_table_data(data: list[list[str]]) -> list[list[str]]:
    """
    Remove colunas vazias/fantasmas para evitar tabelas com dezenas de colunas
    estreitas que quebram o texto em uma letra por linha no DOCX.
    """
    rows = []
    for row in data or []:
        vals = [str(v or "").strip() for v in (row or [])]
        if any(vals):
            rows.append(vals)
    if not rows:
        return []

    n_cols = max(len(r) for r in rows)
    matrix = [r + [""] * (n_cols - len(r)) for r in rows]

    non_empty = [sum(1 for r in matrix if r[c]) for c in range(n_cols)]
    min_non_empty = max(1, int(len(matrix) * 0.15))
    keep = [c for c in range(n_cols) if non_empty[c] >= min_non_empty]
    if not keep:
        return []

    compact = [[r[c] for c in keep] for r in matrix]
    # recorta bordas totalmente vazias que sobraram
    while compact and compact[0] and all((not r[0]) for r in compact):
        compact = [r[1:] for r in compact]
    while compact and compact[0] and all((not r[-1]) for r in compact):
        compact = [r[:-1] for r in compact]

    compact = [r for r in compact if any(r)]
    return compact


def _bbox_iou(a: tuple[float, float, float, float], b: tuple[float, float, float, float]) -> float:
    ax0, ay0, ax1, ay1 = a
    bx0, by0, bx1, by1 = b
    ix0, iy0 = max(ax0, bx0), max(ay0, by0)
    ix1, iy1 = min(ax1, bx1), min(ay1, by1)
    iw, ih = max(0.0, ix1 - ix0), max(0.0, iy1 - iy0)
    inter = iw * ih
    if inter <= 0:
        return 0.0
    area_a = max(0.0, (ax1 - ax0) * (ay1 - ay0))
    area_b = max(0.0, (bx1 - bx0) * (by1 - by0))
    union = area_a + area_b - inter
    return inter / union if union > 0 else 0.0


def _find_tables_robust(plumb_page) -> list[dict]:
    """
    Tenta detectar tabelas com múltiplas estratégias do pdfplumber.
    Retorna tabelas deduplicadas e já ordenadas por Y.
    """
    settings_list = [
        # padrão robusto para tabelas com linhas
        {
            "vertical_strategy": "lines",
            "horizontal_strategy": "lines",
            "intersection_tolerance": 5,
            "snap_tolerance": 3,
            "join_tolerance": 3,
            "edge_min_length": 15,
            "min_words_vertical": 2,
            "min_words_horizontal": 1,
        },
        # útil quando bordas não estão contínuas
        {
            "vertical_strategy": "text",
            "horizontal_strategy": "text",
            "intersection_tolerance": 7,
            "snap_tolerance": 4,
            "join_tolerance": 4,
            "min_words_vertical": 2,
            "min_words_horizontal": 1,
        },
        # híbrido
        {
            "vertical_strategy": "lines",
            "horizontal_strategy": "text",
            "intersection_tolerance": 6,
            "snap_tolerance": 4,
            "join_tolerance": 4,
            "edge_min_length": 12,
            "min_words_vertical": 2,
            "min_words_horizontal": 1,
        },
    ]

    found: list[dict] = []
    for settings in settings_list:
        try:
            tables = plumb_page.find_tables(table_settings=settings)
        except Exception:
            continue
        for tbl in tables:
            try:
                data = tbl.extract()
            except Exception:
                data = None
            if not data:
                continue
            # limpa linhas/células vazias e remove colunas fantasmas
            clean = _compact_table_data(data)
            if not clean:
                continue
            x0, y0, x1, y1 = tuple(float(v) for v in tbl.bbox)
            if (x1 - x0) < 40 or (y1 - y0) < 20:
                continue

            # Geometria opcional (só usa quando bater com colunas após compactação).
            try:
                col_bboxes = [c.bbox for c in tbl.columns]
                col_widths_pt = [max(8.0, float(c[2] - c[0])) for c in col_bboxes]
            except Exception:
                col_widths_pt = []
            if len(col_widths_pt) != max((len(r) for r in clean), default=0):
                col_widths_pt = []
            try:
                row_bboxes = [r.bbox for r in tbl.rows]
                row_heights_pt = [max(8.0, float(r[3] - r[1])) for r in row_bboxes]
            except Exception:
                row_heights_pt = []
            if len(row_heights_pt) != len(clean):
                row_heights_pt = []

            found.append({
                "y": y0,
                "x0": x0,
                "x1": x1,
                "bottom": y1,
                "data": clean,
                "col_widths_pt": col_widths_pt,
                "row_heights_pt": row_heights_pt,
                "bbox": (x0, y0, x1, y1),
            })

    # Dedup por overlap alto de bbox.
    deduped: list[dict] = []
    for t in sorted(found, key=lambda d: (d["y"], d["x0"])):
        if any(_bbox_iou(t["bbox"], k["bbox"]) >= 0.72 for k in deduped):
            continue
        deduped.append(t)

    for t in deduped:
        t.pop("bbox", None)
    return deduped



def _page_text_lines(plumb_page, table_bboxes: list) -> list[dict]:
    """
    Extrai linhas com formatação usando extract_words() — já separa palavras
    corretamente, sem palavras coladas.
    """
    def in_table(w):
        return any(
            w["x0"] >= tb[0] - 2 and w["top"] >= tb[1] - 2 and
            w["x1"] <= tb[2] + 2 and w["bottom"] <= tb[3] + 2
            for tb in table_bboxes
        )

    try:
        words = plumb_page.extract_words(
            extra_attrs=["fontname", "size", "non_stroking_color"],
            keep_blank_chars=False,
            use_text_flow=True,
        )
    except TypeError:
        # versões mais antigas não têm use_text_flow
        words = plumb_page.extract_words(
            extra_attrs=["fontname", "size", "non_stroking_color"],
            keep_blank_chars=False,
        )

    words = [w for w in words if not in_table(w)]

    # Agrupa palavras em linhas (tolerância dinâmica no top)
    line_map: dict[float, list] = {}
    for word in words:
        top = word["top"]
        sz = float(word.get("size", 11) or 11)
        tol = max(2.0, min(5.0, sz * 0.32))
        key = next((k for k in line_map if abs(k - top) <= tol), top)
        line_map.setdefault(key, []).append(word)

    result = []
    for top_key in sorted(line_map):
        line_words = sorted(line_map[top_key], key=lambda w: w["x0"])
        text = " ".join(w["text"] for w in line_words).strip()
        if not text:
            continue
        dom  = max(line_words, key=lambda w: w.get("size", 0))
        font = dom.get("fontname", "").lower()
        x0 = min(w["x0"] for w in line_words)
        x1 = max(w["x1"] for w in line_words)
        tops = [float(w.get("top", top_key)) for w in line_words]
        bottoms = [float(w.get("bottom", top_key + 12)) for w in line_words]
        result.append({
            "y":      top_key,
            "x0":     x0,
            "x1":     x1,
            "text":   text,
            "bold":   "bold" in font,
            "italic": "italic" in font or "oblique" in font,
            "size":   _norm_size(dom.get("size", 11)),
            "color":  _color_to_hex(dom.get("non_stroking_color", 0)),
            "line_height": max(8.0, max(bottoms) - min(tops)),
        })
    return result


def _page_images(pypdf_page, plumb_page, page_w: float) -> list[dict]:
    """Extrai imagens com posição via pypdf (BSD-3) + pdfplumber para bbox."""
    from PIL import Image as PILImage

    plumb_imgs = list(plumb_page.images)

    # Busca exaustiva de posição: tenta todas as variações do nome
    def find_meta(raw_name: str) -> dict:
        candidates = {raw_name, raw_name.lstrip("/"), "/" + raw_name.lstrip("/")}
        for img in plumb_imgs:
            pname = img.get("name", "")
            if pname in candidates or pname.lstrip("/") in candidates:
                return img
        return {}

    # Fallback: lista na ordem natural do PDF (sem reordenar)
    pypdf_list = list(pypdf_page.images)

    imgs = []
    for idx, pypdf_img in enumerate(pypdf_list):
        try:
            raw_name = getattr(pypdf_img, "name", "")
            meta     = find_meta(raw_name)

            # Se ainda não achou por nome, usa índice direto
            if not meta and idx < len(plumb_imgs):
                meta = plumb_imgs[idx]

            pil_img = pypdf_img.image
            if pil_img is None:
                pil_img = PILImage.open(io.BytesIO(pypdf_img.data))
            if pil_img.mode not in ("RGB", "L"):
                pil_img = pil_img.convert("RGB")
            w, h = pil_img.size
            # Permite logos pequenos.
            if w < 6 or h < 6:
                continue

            buf = io.BytesIO()
            pil_img.save(buf, format="PNG")
            buf.seek(0)

            if meta:
                x0    = meta.get("x0", 0)
                x1    = meta.get("x1", x0 + w)
                y_pos = meta.get("top", 0)
                y_bottom = meta.get("bottom", y_pos + h)
                ratio = (x0 + x1) / 2 / page_w if page_w > 0 else 0.0
                has_pos = True
            else:
                # sem posição: assume canto superior esquerdo
                y_pos, ratio = 0.0, 0.0
                x0, x1 = 0.0, float(w)
                y_bottom = float(h)
                has_pos = False

            imgs.append({
                "data": buf,
                "w": w,
                "h": h,
                "y": y_pos,
                "ratio": ratio,
                "x0": x0,
                "x1": x1,
                "bottom": y_bottom,
                "has_pos": has_pos,
            })
        except Exception:
            continue

    imgs.sort(key=lambda x: x["y"])
    return imgs


def convert_pdf_to_xlsx(pdf_path: str, output_path: str) -> None:
    """
    Converte PDF → XLSX preservando imagens, formatação de texto e tabelas.
    Usa apenas bibliotecas MIT/BSD-3 (100% comercial).
    """
    import pdfplumber
    import xlsxwriter
    from pypdf import PdfReader

    reader = PdfReader(pdf_path)
    pages_data = []
    try:
        with pdfplumber.open(pdf_path) as plumber_pdf:
            for pg_idx, plumb_page in enumerate(plumber_pdf.pages):
                page_w = float(plumb_page.width)

                table_bboxes, tables = [], []
                for tbl in sorted(plumb_page.find_tables(), key=lambda t: t.bbox[1]):
                    data = tbl.extract()
                    if data:
                        table_bboxes.append(tbl.bbox)
                        tables.append({"y": tbl.bbox[1], "data": data})

                text_lines = _page_text_lines(plumb_page, table_bboxes)
                images = _page_images(reader.pages[pg_idx], plumb_page, page_w)
                pages_data.append({
                    "num": pg_idx + 1,
                    "images": images,
                    "tables": tables,
                    "text_lines": text_lines,
                })
    finally:
        try:
            reader.close()
        except Exception:
            pass

    workbook = xlsxwriter.Workbook(output_path, {"strings_to_urls": False})
    try:
        ws = workbook.add_worksheet("Documento")
        _fmt_cache: dict = {}

        def get_fmt(**kw):
            key = tuple(sorted(kw.items()))
            if key not in _fmt_cache:
                _fmt_cache[key] = workbook.add_format(kw)
            return _fmt_cache[key]

        def text_fmt(bold, italic, size, color):
            return get_fmt(
                bold=bold, italic=italic, font_size=size, font_color=color, text_wrap=True, valign="vcenter"
            )

        fmt_sep = get_fmt(bold=True, font_size=10, font_color="#FFFFFF", bg_color="#1E3A5F")
        fmt_th = get_fmt(
            bold=True, bg_color="#1E3A5F", font_color="#FFFFFF", border=1, text_wrap=True,
            align="center", valign="vcenter"
        )
        fmt_td = get_fmt(border=1, text_wrap=True, valign="vcenter")
        fmt_td_alt = get_fmt(border=1, bg_color="#EBF2FA", text_wrap=True, valign="vcenter")

        max_w, max_h = 300, 180
        sheet_px_w = 700
        cur_row = 0
        max_col = 0
        col_widths: dict[int, int] = {}

        def track(col, val):
            col_widths[col] = max(col_widths.get(col, 8), min(len(str(val)) + 2, 55))

        for page in pages_data:
            if cur_row > 0:
                cur_row += 1

            if page["images"]:
                img_rows = 0
                for img in page["images"]:
                    scale = min(max_w / img["w"], max_h / img["h"], 1.0)
                    rows = max(1, int(img["h"] * scale / 15) + 1)
                    img_rows = max(img_rows, rows)
                    for r in range(cur_row, cur_row + rows):
                        ws.set_row(r, 15)

                    ratio = img.get("ratio", 0.0)
                    x_offset = int(ratio * sheet_px_w)
                    scaled_w = img["w"] * scale
                    x_offset = max(0, min(x_offset, int(sheet_px_w - scaled_w)))

                    img["data"].seek(0)
                    ws.insert_image(cur_row, 0, "img.png", {
                        "image_data": img["data"],
                        "x_scale": scale,
                        "y_scale": scale,
                        "x_offset": x_offset,
                        "object_position": 1,
                    })
                cur_row += img_rows + 1

            ws.write(cur_row, 0, f"Página {page['num']}", fmt_sep)
            cur_row += 1

            events = (
                [("text", t["y"], t) for t in page["text_lines"]] +
                [("table", t["y"], t) for t in page["tables"]]
            )
            events.sort(key=lambda e: e[1])

            for kind, _, content in events:
                if kind == "text":
                    txt = content["text"]
                    ws.write(cur_row, 0, txt, text_fmt(content["bold"], content["italic"], content["size"], content["color"]))
                    track(0, txt)
                    cur_row += 1
                    continue

                n_cols = max((len(r) for r in content["data"] if r), default=1)
                max_col = max(max_col, n_cols - 1)
                for r_idx, trow in enumerate(content["data"]):
                    if not trow:
                        continue
                    fmt = fmt_th if r_idx == 0 else (fmt_td_alt if r_idx % 2 == 0 else fmt_td)
                    for c_idx, val in enumerate(trow):
                        v = str(val or "").strip()
                        ws.write(cur_row, c_idx, v, fmt)
                        track(c_idx, v)
                    cur_row += 1
                cur_row += 1

        for c in range(max(max_col, 0) + 1):
            ws.set_column(c, c, col_widths.get(c, 14))
    finally:
        workbook.close()


def ocr_pdf_to_docx(pdf_path: str, output_path: str) -> None:
    """
    OCR para PDFs escaneados → DOCX via Tesseract + pypdfium2.
    Licença pypdfium2: Apache-2.0 — uso comercial livre, sem PyMuPDF/AGPL.
    """
    try:
        import re
        import pytesseract
        import pypdfium2 as pdfium
        from docx import Document
        from docx.shared import Pt
    except Exception as import_err:
        logger.warning("Dependências de OCR indisponíveis (%s). Fallback para conversão normal.", import_err)
        convert_pdf_to_docx(pdf_path, output_path)
        return

    _configure_tesseract(pytesseract)
    try:
        pytesseract.get_tesseract_version()
    except Exception as exc:
        logger.warning("Tesseract indisponível (%s). Fazendo fallback para conversão normal.", exc)
        convert_pdf_to_docx(pdf_path, output_path)
        return

    def _norm_text(txt: str) -> str:
        txt = (txt or "").strip()
        txt = re.sub(r"\s+", " ", txt)
        # reduz ruído de OCR (símbolos repetidos)
        txt = re.sub(r"([^\w\s])\1{2,}", r"\1", txt, flags=re.UNICODE)
        return txt

    def _ocr_words(img, psm: int = 6, min_conf: float = 55.0):
        cfg = f"--oem 1 --psm {psm}"
        for lang in ("por+eng", "eng"):
            try:
                data = pytesseract.image_to_data(
                    img, lang=lang, config=cfg, output_type=pytesseract.Output.DICT
                )
                words = []
                n = len(data.get("text", []))
                for k in range(n):
                    txt = _norm_text(data["text"][k])
                    if not txt:
                        continue
                    try:
                        conf = float(data["conf"][k])
                    except Exception:
                        conf = 0.0
                    if conf < min_conf:
                        continue
                    words.append({
                        "text": txt,
                        "conf": conf,
                        "left": int(data["left"][k]),
                        "top": int(data["top"][k]),
                        "block_num": int(data["block_num"][k]),
                        "par_num": int(data["par_num"][k]),
                        "line_num": int(data["line_num"][k]),
                    })
                if words:
                    return words
            except Exception:
                continue
        return []

    def _words_to_lines(words):
        if not words:
            return []
        line_map: dict[tuple[int, int, int], list] = {}
        for w in words:
            key = (w["block_num"], w["par_num"], w["line_num"])
            line_map.setdefault(key, []).append(w)
        lines = []
        for grp in line_map.values():
            grp = sorted(grp, key=lambda x: x["left"])
            txt = _norm_text(" ".join(x["text"] for x in grp))
            if len(txt) < 2:
                continue
            avg_conf = sum(x["conf"] for x in grp) / len(grp)
            if avg_conf < 60:
                continue
            lines.append({"top": min(x["top"] for x in grp), "text": txt})
        lines.sort(key=lambda x: x["top"])
        return lines

    def _ocr_cell_text(img_crop):
        words = _ocr_words(img_crop, psm=7, min_conf=50.0)
        if not words:
            return ""
        words = sorted(words, key=lambda x: (x["top"], x["left"]))
        return _norm_text(" ".join(w["text"] for w in words))

    def _extract_tables(img):
        """
        Detecção conservadora de tabelas com grade.
        Só retorna tabela quando há linhas/células claras, para evitar falsos positivos.
        """
        try:
            import cv2
            import numpy as np
        except Exception:
            return []

        gray = np.array(img.convert("L"))
        h_img, w_img = gray.shape[:2]
        bw = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 31, 12
        )
        h_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (max(24, w_img // 30), 1))
        v_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, max(24, h_img // 30)))
        h_lines = cv2.morphologyEx(bw, cv2.MORPH_OPEN, h_kernel, iterations=1)
        v_lines = cv2.morphologyEx(bw, cv2.MORPH_OPEN, v_kernel, iterations=1)
        grid = cv2.add(h_lines, v_lines)

        contours, _ = cv2.findContours(grid, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        tables = []
        for c in sorted(contours, key=lambda c0: cv2.boundingRect(c0)[1]):
            x, y, w, h = cv2.boundingRect(c)
            if w < max(180, int(w_img * 0.30)) or h < 70:
                continue
            if (w * h) < int((w_img * h_img) * 0.012):
                continue

            roi = grid[y:y + h, x:x + w]
            cell_contours, _ = cv2.findContours(roi, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
            cell_boxes = []
            for cc in cell_contours:
                cx, cy, cw, ch = cv2.boundingRect(cc)
                if cw < 30 or ch < 16:
                    continue
                if cw > int(w * 0.98) and ch > int(h * 0.98):
                    continue
                cell_boxes.append((x + cx, y + cy, cw, ch))

            if len(cell_boxes) < 4:
                continue

            rows: list[list[tuple[int, int, int, int]]] = []
            for bx in sorted(cell_boxes, key=lambda b: (b[1], b[0])):
                if not rows or abs(rows[-1][0][1] - bx[1]) > 10:
                    rows.append([bx])
                else:
                    rows[-1].append(bx)
            if len(rows) < 2:
                continue

            text_rows = []
            for row in rows:
                row = sorted(row, key=lambda b: b[0])
                vals = []
                for cx, cy, cw, ch in row:
                    pad = 2
                    crop = img.crop((max(0, cx + pad), max(0, cy + pad), cx + cw - pad, cy + ch - pad))
                    vals.append(_ocr_cell_text(crop))
                if any(v.strip() for v in vals):
                    text_rows.append(vals)

            if len(text_rows) >= 2:
                tables.append({"top": y, "bottom": y + h, "rows": text_rows})

        return tables

    doc = Document()
    pdf = None
    try:
        pdf = pdfium.PdfDocument(pdf_path)
    except Exception as open_err:
        logger.warning("Falha ao abrir PDF no OCR (%s). Fallback para conversão normal.", open_err)
        convert_pdf_to_docx(pdf_path, output_path)
        return

    try:
        for i, page in enumerate(pdf):
            try:
                bitmap = page.render(scale=300 / 72)
                img = bitmap.to_pil()
                if i > 0:
                    doc.add_page_break()

                tables = _extract_tables(img)
                words = _ocr_words(img, psm=6, min_conf=55.0)
                lines = _words_to_lines(words)
                if not lines and not tables:
                    doc.add_paragraph("[Sem texto OCR detectado nesta página]")
                    continue

                # Remove linhas dentro de áreas de tabela.
                lines = [
                    ln for ln in lines
                    if not any((tb["top"] - 4) <= ln["top"] <= (tb["bottom"] + 4) for tb in tables)
                ]
                events = [("line", ln["top"], ln) for ln in lines] + [("table", tb["top"], tb) for tb in tables]
                events.sort(key=lambda e: e[1])

                for kind, _, payload in events:
                    if kind == "line":
                        p = doc.add_paragraph(payload["text"])
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        if p.runs:
                            p.runs[0].font.size = Pt(11)
                        continue

                    rows = payload["rows"]
                    cols = max((len(r) for r in rows), default=1)
                    if cols <= 0:
                        continue
                    tbl = doc.add_table(rows=len(rows), cols=cols)
                    tbl.style = "Table Grid"
                    for r_idx, row in enumerate(rows):
                        for c_idx in range(cols):
                            val = row[c_idx].strip() if c_idx < len(row) else ""
                            cell = tbl.cell(r_idx, c_idx)
                            cell.text = val
                            if r_idx == 0 and cell.paragraphs and cell.paragraphs[0].runs:
                                cell.paragraphs[0].runs[0].bold = True
                    doc.add_paragraph("")
            except Exception as page_err:
                logger.warning("Falha no OCR da página %s (%s).", i + 1, page_err)
                if i > 0:
                    doc.add_page_break()
                doc.add_paragraph(f"[Falha ao processar OCR da página {i + 1}]")
        doc.save(output_path)
    finally:
        try:
            if pdf is not None:
                pdf.close()
        except Exception:
            pass


def _create_emergency_docx_from_pdf(pdf_path: str, output_path: str) -> None:
    """
    Último fallback: garante um DOCX válido mesmo em falhas de OCR/conversores.
    """
    from docx import Document
    from pypdf import PdfReader

    doc = Document()
    reader = PdfReader(pdf_path)
    try:
        for i, pg in enumerate(reader.pages):
            if i > 0:
                doc.add_page_break()
            txt = (pg.extract_text() or "").strip()
            if txt:
                doc.add_paragraph(txt)
            else:
                doc.add_paragraph(f"Página {i + 1} (sem texto detectado)")
        doc.save(output_path)
    finally:
        try:
            reader.close()
        except Exception:
            pass


def _should_use_ocr(pdf_path: str) -> bool:
    """
    Decide se OCR realmente deve ser usado.
    Se o PDF já tem texto nativo em boa quantidade, OCR é evitado para não degradar resultado.
    """
    from pypdf import PdfReader

    reader = PdfReader(pdf_path)
    try:
        page_count = len(reader.pages)
        sample = min(page_count, 8)
        if sample == 0:
            return True

        total_chars = 0
        text_pages = 0
        for i in range(sample):
            txt = (reader.pages[i].extract_text() or "").strip()
            c = len(txt)
            total_chars += c
            if c >= 20:
                text_pages += 1

        # Heurística conservadora:
        # se há texto em >= 40% das páginas amostradas ou muitos caracteres totais,
        # tratamos como PDF digital (não usar OCR).
        if text_pages >= max(1, int(sample * 0.4)) or total_chars >= 300:
            return False
        return True
    finally:
        try:
            reader.close()
        except Exception:
            pass


def convert_pdf_to_docx_with_ocr_fallback(pdf_path: str, output_path: str) -> None:
    """
    OCR resiliente: tenta OCR, depois conversão normal, e por fim DOCX de emergência.
    """
    try:
        ocr_pdf_to_docx(pdf_path, output_path)
        if os.path.isfile(output_path) and os.path.getsize(output_path) > 0:
            return
        raise RuntimeError("OCR não gerou arquivo de saída válido.")
    except Exception as ocr_err:
        logger.warning("OCR falhou (%s). Tentando conversão DOCX normal.", ocr_err)

    try:
        convert_pdf_to_docx(pdf_path, output_path)
        if os.path.isfile(output_path) and os.path.getsize(output_path) > 0:
            return
        raise RuntimeError("Conversão DOCX padrão não gerou arquivo de saída válido.")
    except Exception as std_err:
        logger.warning("Conversão DOCX padrão falhou (%s). Gerando DOCX de emergência.", std_err)

    _create_emergency_docx_from_pdf(pdf_path, output_path)


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/convert", methods=["POST"])
def convert():
    if "file" not in request.files:
        return jsonify({"error": "Nenhum arquivo enviado."}), 400

    file = request.files["file"]
    output_format = request.form.get("format", "docx").lower()
    use_ocr = request.form.get("ocr", "false").lower() == "true"

    if file.filename == "":
        return jsonify({"error": "Nome de arquivo inválido."}), 400

    if not allowed_file(file.filename):
        return jsonify({"error": "Apenas arquivos PDF são aceitos."}), 400

    if output_format not in ("docx", "odt", "xlsx"):
        return jsonify({"error": "Formato inválido. Use 'docx', 'odt' ou 'xlsx'."}), 400

    original_name = Path(secure_filename(file.filename)).stem
    output_filename = f"{original_name}.{output_format}"

    input_path = None
    output_path = None
    try:
        base_tmp = _pick_runtime_tmpdir()
        input_path = os.path.join(base_tmp, f"conversor-{uuid.uuid4()}.pdf")
        output_path = os.path.join(base_tmp, f"conversor-{uuid.uuid4()}.{output_format}")

        file.save(input_path)

        if output_format == "docx":
            if use_ocr:
                logger.info("OCR forçado pelo usuário.")
                convert_pdf_to_docx_with_ocr_fallback(input_path, output_path)
            else:
                # Modo estrito (sem OCR): pipeline fixo PDF -> ODT -> DOCX.
                try:
                    intermediate_odt = os.path.join(base_tmp, f"conversor-{uuid.uuid4()}.odt")
                    _odt_via_libreoffice(input_path, intermediate_odt)
                    try:
                        _odt_to_docx_via_word_com(intermediate_odt, output_path)
                    except Exception as word_err:
                        logger.warning("ODT->DOCX via Word COM falhou (%s). Tentando via LibreOffice.", word_err)
                        _odt_to_docx_via_libreoffice(intermediate_odt, output_path)
                    try:
                        if os.path.exists(intermediate_odt):
                            os.remove(intermediate_odt)
                    except Exception:
                        pass
                except Exception as docx_err:
                    logger.warning("Conversão DOCX (PDF->ODT->DOCX) falhou: %s", docx_err)
                    short_reason = str(docx_err).strip()
                    if len(short_reason) > 220:
                        short_reason = short_reason[:220].rstrip() + "..."
                    return jsonify({
                        "error": (
                            "Não foi possível converter este PDF para DOCX pelo fluxo ODT. "
                            f"Detalhe: {short_reason}"
                        )
                    }), 422
        elif output_format == "odt":
            try:
                _odt_via_libreoffice(input_path, output_path)
            except Exception as odt_err:
                logger.warning("Conversão ODT falhou no LibreOffice: %s", odt_err)
                return jsonify({
                    "error": (
                        "Não foi possível converter este PDF para ODT neste ambiente. "
                        "Tente DOCX com OCR ou verifique a instalação do LibreOffice."
                    )
                }), 422
        else:
            convert_pdf_to_xlsx(input_path, output_path)

        if not os.path.isfile(output_path) or os.path.getsize(output_path) == 0:
            logger.error("Conversão finalizou sem gerar arquivo de saída válido")
            return jsonify({"error": "Falha na conversão: arquivo de saída inválido."}), 500

        mime = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if output_format == "docx"
            else (
                "application/vnd.oasis.opendocument.text"
                if output_format == "odt"
                else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        )

        with open(output_path, "rb") as f:
            data = f.read()
    except Exception as e:
        logger.exception("Erro na conversão")
        return jsonify({"error": f"Falha na conversão: {str(e)}"}), 500
    finally:
        if input_path and os.path.exists(input_path):
            try:
                os.remove(input_path)
            except Exception:
                pass
        if output_path and os.path.exists(output_path):
            try:
                os.remove(output_path)
            except Exception:
                pass

    return send_file(
        io.BytesIO(data),
        mimetype=mime,
        as_attachment=True,
        download_name=output_filename,
    )


@app.route("/health")
def health():
    return jsonify({"status": "ok", "app": "CamargoApps PDF Converter"})


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
